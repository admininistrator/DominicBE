import io, sys
sys.path.insert(0, '.')

with open('scripts/test_assignments.docx', 'rb') as f:
    content = f.read()

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

doc = Document(io.BytesIO(content))
print('Paragraphs:', len(doc.paragraphs))
for p in doc.paragraphs[:15]:
    if p.text.strip():
        print('  P:', repr(p.text[:120]))

print('Tables:', len(doc.tables))
for i, t in enumerate(doc.tables[:5]):
    for row in t.rows[:5]:
        for cell in row.cells[:5]:
            if cell.text.strip():
                print(f'  T{i}:', repr(cell.text[:120]))

# Check for text boxes / drawing objects
body_xml = doc.element.body.xml
import re
text_tags = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', body_xml)
print(f'\nAll <w:t> text nodes ({len(text_tags)}):')
for t in text_tags[:30]:
    if t.strip():
        print(f'  {repr(t[:100])}')

# Check headers/footers
print('\nHeaders/footers:')
for section in doc.sections:
    for hdr in [section.header, section.footer]:
        for p in hdr.paragraphs:
            if p.text.strip():
                print(' ', repr(p.text[:80]))

