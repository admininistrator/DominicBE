import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from docx import Document
from lxml import etree

doc = Document(os.path.join(os.path.dirname(__file__), 'test_assignments.docx'))
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Check row 0 "Thong tin"
row = doc.tables[0].rows[0]
cell = row.cells[0]
print("Row 0 cell 0 text repr:", repr(cell.text))
print("chars:")
for ch in cell.text[:30]:
    print(f"  U+{ord(ch):04X} = {repr(ch)}")

# Check row 5
row5 = doc.tables[0].rows[5]
cell5 = row5.cells[0]
first_text = next(cell5._element.iter(f"{{{WORD_NS}}}t")).text or ""
print("\nRow5 first w:t repr:", repr(first_text))
for ch in first_text:
    print(f"  U+{ord(ch):04X} = {repr(ch)}")
